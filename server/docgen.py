"""
Document generation for bets: executive Memo + PRD, as .docx and .pdf.

One Gemini call (made by pyserver) returns structured JSON content; this module
turns (kind, content, bet) into a neutral block list and renders it twice —
python-docx for the editable .docx, reportlab for the inline-previewable .pdf —
so the visual treatment lives in exactly one place per format.

Templates are modelled on real Astra docs: the Botim Pots PRD (meta block,
"In one line", narrative sections, grouped requirements, success-metrics table
with a decision gate) and the Astra InsurTech executive memo (summary →
opportunity → market → model → financials → risks → recommendation).
"""

import io
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# ---------------------------------------------------------------- palette
ACCENT = "4F46E5"      # indigo — matches the app's primary
INK = "1E293B"         # dark slate body/headings
MUTED = "64748B"       # secondary text
HEADER_FILL = "EEF2FF" # table header band
BORDER = "CBD5E1"

# ---------------------------------------------------------------- Gemini contract

PRD_SCHEMA = {
    "type": "object",
    "required": [
        "oneLiner", "objective", "whyItWorks", "whoPrimary", "whoSecondary",
        "whoNotTarget", "principles", "requirements", "outOfScope",
        "metrics", "decisionGate",
    ],
    "properties": {
        "oneLiner": {"type": "string"},
        "objective": {"type": "array", "items": {"type": "string"}},
        "whyItWorks": {"type": "array", "items": {"type": "string"}},
        "whoPrimary": {"type": "string"},
        "whoSecondary": {"type": "string"},
        "whoNotTarget": {"type": "string"},
        "principles": {"type": "array", "items": {"type": "string"}},
        "requirements": {"type": "array", "items": {
            "type": "object", "required": ["heading", "items"],
            "properties": {
                "heading": {"type": "string"},
                "items": {"type": "array", "items": {"type": "string"}},
            }}},
        "outOfScope": {"type": "array", "items": {"type": "string"}},
        "metrics": {"type": "array", "items": {
            "type": "object", "required": ["metric", "definition", "target"],
            "properties": {
                "metric": {"type": "string"},
                "definition": {"type": "string"},
                "target": {"type": "string"},
            }}},
        "decisionGate": {"type": "string"},
    },
}

MEMO_SCHEMA = {
    "type": "object",
    "required": [
        "executiveSummary", "opportunity", "model",
        "financialOutlook", "recommendation", "nextSteps",
    ],
    "properties": {
        "executiveSummary": {"type": "array", "items": {"type": "string"}},
        "opportunity": {"type": "array", "items": {"type": "string"}},
        "model": {"type": "array", "items": {"type": "string"}},
        "financialOutlook": {"type": "array", "items": {"type": "string"}},
        "recommendation": {"type": "array", "items": {"type": "string"}},
        "nextSteps": {"type": "array", "items": {"type": "string"}},
    },
}


DEPTH_RULES = (
    "Depth requirement — this document must be MORE specific than the dashboard, never a "
    "restatement of it:\n"
    "- Do not copy sentences from the bet's description or aiSummary; re-derive every section "
    "from the underlying data.\n"
    "- Translate KPI values into operational implications (e.g. a 46% activation rate means "
    "X of Y pilot users completed the core action — say what that action is and what moves it).\n"
    "- Where the bet gives TAM/SAM/SOM and KPI values, do the arithmetic: revenue per user, "
    "implied year-1 volumes, what share of SOM the targets represent.\n"
    "- Name concrete mechanisms, flows, limits, SLAs, and edge cases — not adjectives. "
    "'Reduce CAC via embedded distribution' is too vague; 'surface the offer on the remittance "
    "confirmation screen, capped at one prompt per week' is the bar.\n"
    "- Where the bet data is silent, reason to a specific, plausible position and mark it '(assumption)'.\n"
)


def prd_prompt(bet, kpi_lines):
    return (
        "You are drafting a Product Requirements Document for Astra Tech (UAE/MENA fintech), "
        "in the house style: confident, concrete, no hedging, written for a PM + designer audience "
        "who will build directly from this document. "
        "Use ONLY the bet data below — do not invent market figures.\n\n"
        + DEPTH_RULES + "\n"
        f"Bet (full JSON):\n{bet}\n\n"
        f"Stage KPI thresholds for context:\n{kpi_lines}\n\n"
        "Produce JSON with:\n"
        "- oneLiner: one sentence stating what the product lets users do and where it lives.\n"
        "- objective: 2-3 paragraphs — the user experience end-to-end and the core mechanic that makes it work.\n"
        "- whyItWorks: 2-3 paragraphs — the behavioural/economic foundation and any secondary engagement effect.\n"
        "- whoPrimary / whoSecondary / whoNotTarget: one paragraph each. whoNotTarget names who we are NOT building for and why.\n"
        "- principles: 3 product principles to protect, each one sentence, phrased as a mental model the design must reinforce.\n"
        "- requirements: 4-7 grouped requirement areas (heading + 2-6 bullet items each), drafted from the description and initiatives. "
        "Each bullet must be buildable and testable: name the screen/state/flow, the limit or threshold with a number, "
        "and the failure/edge case behaviour. Mark unknowns '(TBD)'.\n"
        "- outOfScope: 3-5 things explicitly NOT in V1.\n"
        "- metrics: one row per KPI the bet tracks — metric (the KPI label), definition (one line), target "
        "(derive from the prioritise band of its threshold, e.g. '>40%'). Include custom KPIs if present.\n"
        "- decisionGate: one sentence naming which metric(s) must hit to proceed.\n"
        "Plain text only — no markdown syntax anywhere."
    )


def memo_prompt(bet, kpi_lines):
    return (
        "You are drafting an executive memo for Astra Tech's leadership (UAE/MENA fintech) "
        "recommending a course of action on a product bet. House style: lead with the numbers "
        "and the ask, short declarative sentences, no hedging. Use ONLY the bet data below — "
        "do not invent figures; cite the bet's own TAM/SAM/SOM and KPI values where relevant.\n\n"
        + DEPTH_RULES + "\n"
        f"Bet (full JSON):\n{bet}\n\n"
        f"Stage KPI thresholds for context:\n{kpi_lines}\n\n"
        "Produce JSON with:\n"
        "- executiveSummary: 2-3 paragraphs. First paragraph = the opportunity in one breath with the "
        "headline numbers and the recommendation. The reader should be able to stop here.\n"
        "- opportunity: 2-3 paragraphs — the problem, why now, why Astra is positioned to win.\n"
        "- model: 2-3 paragraphs — proposed product/partnership model, distribution, and what it costs us to run.\n"
        "- financialOutlook: 1-2 paragraphs reading the bet's current KPI values against the stage thresholds "
        "and what they imply about unit economics.\n"
        "- recommendation: 1-2 paragraphs — explicit call aligned with the bet's decision and score.\n"
        "- nextSteps: 3-5 concrete actions with owners or timeframes where known.\n"
        "Plain text only — no markdown syntax anywhere."
    )


def revise_prompt(kind, current_json, instructions, bet, kpi_lines):
    label = "PRD" if kind == "prd" else "executive memo"
    return (
        f"You previously drafted the {label} content below (structured JSON) for an Astra Tech "
        "product bet. The user has requested a change. Apply it and return the FULL updated JSON "
        "in the exact same schema — keep every section you were not asked to change as-is, except "
        "where the change makes them inconsistent.\n\n"
        + DEPTH_RULES + "\n"
        f"Current document content:\n{current_json}\n\n"
        f"Requested change: {instructions}\n\n"
        f"Bet data for reference:\n{bet}\n\n"
        f"Stage KPI thresholds:\n{kpi_lines}\n"
        "Plain text only — no markdown syntax anywhere."
    )


# ---------------------------------------------------------------- bet data → display helpers

def _fmt_kpi(defn, v):
    if v is None or v == "":
        return "—"
    f = defn.get("format")
    try:
        if f == "pct":
            return f"{round(float(v) * 100)}%"
        if f == "x":
            return f"{float(v):.1f}x"
        if f == "months":
            return f"{v} months"
        if f == "weeks":
            return f"{v} weeks"
    except (TypeError, ValueError):
        pass
    return str(v)


def _kpi_rows(bet, kpi_defs):
    hidden = set(bet.get("hiddenKpis") or [])
    rows = []
    for k, d in kpi_defs.items():
        if k in hidden:
            continue
        rows.append([d["label"], _fmt_kpi(d, (bet.get("kpis") or {}).get(k)), d["thresholds"]])
    for c in bet.get("customKpis") or []:
        bands = f"{c.get('kill', '?')} kill · {c.get('proceed', '?')} proceed · {c.get('prioritise', '?')} prioritise"
        rows.append([c.get("name", ""), str(c.get("value") or "—"), bands])
    return rows


# ---------------------------------------------------------------- neutral block model
# block kinds: kicker, title, meta(rows), h2, h3, p, bullets(items), table(headers, rows), gate

def build_blocks(kind, content, bet, kpi_defs):
    today = datetime.now(timezone.utc).strftime("%d %b %Y")
    score = bet.get("score")
    meta = [
        ("Owner", "Astra Tech · Strategy"),
        ("Stage", f"{bet.get('stage', '—')} · {bet.get('decision', '—')}"),
        ("Score", f"{score} / 100" if score is not None else "—"),
        ("Date", today),
    ]
    blocks = []

    if kind == "prd":
        blocks += [
            ("kicker", "Product Requirements Document"),
            ("title", bet.get("name", "")),
            ("meta", meta),
            ("h2", "In one line"),
            ("p", content["oneLiner"]),
            ("h2", "What we're trying to achieve"),
        ]
        blocks += [("p", t) for t in content["objective"]]
        blocks.append(("h2", "Why we think this works"))
        blocks += [("p", t) for t in content["whyItWorks"]]
        blocks += [
            ("h2", "Who it's for"),
            ("h3", "Primary"),
            ("p", content["whoPrimary"]),
            ("h3", "Secondary"),
            ("p", content["whoSecondary"]),
            ("h3", "Not the target"),
            ("p", content["whoNotTarget"]),
            ("h2", "Principles to protect"),
            ("bullets", content["principles"]),
            ("h2", "Requirements"),
        ]
        for group in content["requirements"]:
            blocks.append(("h3", group["heading"]))
            blocks.append(("bullets", group["items"]))
        blocks += [
            ("h2", "Explicitly out of scope for V1"),
            ("bullets", content["outOfScope"]),
            ("h2", "Success metrics"),
            ("table", (["Metric", "Definition", "Target"],
                       [[m["metric"], m["definition"], m["target"]] for m in content["metrics"]])),
            ("gate", f"Decision gate: {content['decisionGate']}"),
        ]

    else:  # memo
        blocks += [
            ("kicker", "Executive Memo"),
            ("title", bet.get("name", "")),
            ("meta", meta),
            ("h2", "Executive summary"),
        ]
        blocks += [("p", t) for t in content["executiveSummary"]]
        blocks.append(("h2", "Opportunity"))
        blocks += [("p", t) for t in content["opportunity"]]

        blocks.append(("h2", "Market & competition"))
        m = bet.get("market") or {}
        sources = m.get("sources") or {}
        blocks.append(("table", (["Measure", "Estimate", "Source"], [
            ["TAM", m.get("tam", "—"), sources.get("tam", "")],
            ["SAM", m.get("sam", "—"), sources.get("sam", "")],
            ["SOM", m.get("som", "—"), sources.get("som", "")],
        ])))
        comps = m.get("competitors") or []
        if comps:
            blocks.append(("table", (["Competitor", "Threat", "Their edge", "Their gap"], [
                [c.get("name", ""), c.get("threat", ""),
                 c.get("edge") or c.get("strength") or "",
                 c.get("gap") or c.get("weakness") or ""]
                for c in comps
            ])))

        blocks.append(("h2", "Proposed model & distribution"))
        blocks += [("p", t) for t in content["model"]]

        blocks.append(("h2", "Financial & KPI outlook"))
        blocks += [("p", t) for t in content["financialOutlook"]]
        kpi_rows = _kpi_rows(bet, kpi_defs)
        if kpi_rows:
            blocks.append(("table", (["KPI", "Current", "Threshold bands"], kpi_rows)))

        risks = bet.get("risks") or []
        if risks:
            blocks.append(("h2", "Risks & mitigations"))
            blocks.append(("table", (["Risk", "Category", "Severity", "Mitigation"], [
                [r.get("name", ""), r.get("category", ""), r.get("severity", ""), r.get("mitigation", "")]
                for r in risks
            ])))

        blocks.append(("h2", "Recommendation & next steps"))
        blocks += [("p", t) for t in content["recommendation"]]
        blocks.append(("bullets", content["nextSteps"]))

    blocks.append(("footer", f"Generated by Alchemy · New Horizons portfolio · {today}"))
    return blocks


# ---------------------------------------------------------------- DOCX renderer

def _docx_shade(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _docx_table(doc, headers, rows, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _docx_shade(cell, HEADER_FILL)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(font_size)
                r.font.color.rgb = RGBColor.from_string(ACCENT)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
                    r.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def build_docx(kind, content, bet, kpi_defs):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Helvetica"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for block in build_blocks(kind, content, bet, kpi_defs):
        tag, payload = block[0], block[1]
        if tag == "kicker":
            p = doc.add_paragraph()
            r = p.add_run(payload.upper())
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(ACCENT)
            p.paragraph_format.space_after = Pt(2)
        elif tag == "title":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.font.size = Pt(24)
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(INK)
            p.paragraph_format.space_after = Pt(12)
        elif tag == "meta":
            table = doc.add_table(rows=2, cols=len(payload))
            table.style = "Table Grid"
            for i, (label, value) in enumerate(payload):
                head = table.rows[0].cells[i]
                head.text = label.upper()
                _docx_shade(head, HEADER_FILL)
                for pp in head.paragraphs:
                    for rr in pp.runs:
                        rr.font.size = Pt(8)
                        rr.font.bold = True
                        rr.font.color.rgb = RGBColor.from_string(MUTED)
                body = table.rows[1].cells[i]
                body.text = value
                for pp in body.paragraphs:
                    for rr in pp.runs:
                        rr.font.size = Pt(9)
                        rr.font.color.rgb = RGBColor.from_string(INK)
            doc.add_paragraph()
        elif tag == "h2":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(ACCENT)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif tag == "h3":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(INK)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
        elif tag == "p":
            doc.add_paragraph(payload)
        elif tag == "bullets":
            for item in payload:
                p = doc.add_paragraph(item, style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
        elif tag == "table":
            headers, rows = payload
            _docx_table(doc, headers, rows)
        elif tag == "gate":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor.from_string(ACCENT)
        elif tag == "footer":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor.from_string(MUTED)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------- PDF renderer

def _pdf_styles():
    ink = colors.HexColor(f"#{INK}")
    accent = colors.HexColor(f"#{ACCENT}")
    muted = colors.HexColor(f"#{MUTED}")
    return {
        "kicker": ParagraphStyle("kicker", fontName="Helvetica-Bold", fontSize=8.5,
                                 textColor=accent, spaceAfter=4, leading=11),
        "title": ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=22,
                                textColor=ink, spaceAfter=12, leading=26),
        "h2": ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=13,
                             textColor=accent, spaceBefore=14, spaceAfter=6, leading=16),
        "h3": ParagraphStyle("h3", fontName="Helvetica-Bold", fontSize=10.5,
                             textColor=ink, spaceBefore=8, spaceAfter=3, leading=13),
        "p": ParagraphStyle("p", fontName="Helvetica", fontSize=9.5,
                            textColor=ink, spaceAfter=7, leading=14, alignment=TA_LEFT),
        "bullet": ParagraphStyle("bullet", fontName="Helvetica", fontSize=9.5,
                                 textColor=ink, spaceAfter=3, leading=13,
                                 leftIndent=14, bulletIndent=4),
        "cell": ParagraphStyle("cell", fontName="Helvetica", fontSize=8.5,
                               textColor=ink, leading=11),
        "cellHead": ParagraphStyle("cellHead", fontName="Helvetica-Bold", fontSize=8.5,
                                   textColor=accent, leading=11),
        "metaHead": ParagraphStyle("metaHead", fontName="Helvetica-Bold", fontSize=7,
                                   textColor=muted, leading=9),
        "gate": ParagraphStyle("gate", fontName="Helvetica-Bold", fontSize=9.5,
                               textColor=accent, spaceBefore=4, leading=13),
        "footer": ParagraphStyle("footer", fontName="Helvetica", fontSize=7.5,
                                 textColor=muted, spaceBefore=18, alignment=1),
    }


def _esc(s):
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_table(headers, rows, styles, col_widths=None):
    head = [Paragraph(_esc(h), styles["cellHead"]) for h in headers]
    body = [[Paragraph(_esc(v), styles["cell"]) for v in row] for row in rows]
    t = Table([head] + body, colWidths=col_widths, repeatRows=1, hAlign="LEFT")
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{HEADER_FILL}")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{BORDER}")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    return t


def build_pdf(kind, content, bet, kpi_defs):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2.0 * cm, bottomMargin=2.0 * cm,
        leftMargin=2.2 * cm, rightMargin=2.2 * cm,
        title=f"{bet.get('name', '')} — {'PRD' if kind == 'prd' else 'Executive Memo'}",
        author="Alchemy · Astra Tech",
    )
    styles = _pdf_styles()
    width = A4[0] - doc.leftMargin - doc.rightMargin
    story = []

    for block in build_blocks(kind, content, bet, kpi_defs):
        tag, payload = block[0], block[1]
        if tag == "kicker":
            story.append(Paragraph(_esc(payload.upper()), styles["kicker"]))
        elif tag == "title":
            story.append(Paragraph(_esc(payload), styles["title"]))
            story.append(HRFlowable(width="100%", thickness=1.2,
                                    color=colors.HexColor(f"#{ACCENT}"), spaceAfter=10))
        elif tag == "meta":
            head = [Paragraph(_esc(label.upper()), styles["metaHead"]) for label, _ in payload]
            vals = [Paragraph(_esc(v), styles["cell"]) for _, v in payload]
            t = Table([head, vals], colWidths=[width / len(payload)] * len(payload), hAlign="LEFT")
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{HEADER_FILL}")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{BORDER}")),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 10))
        elif tag == "h2":
            story.append(Paragraph(_esc(payload), styles["h2"]))
        elif tag == "h3":
            story.append(Paragraph(_esc(payload), styles["h3"]))
        elif tag == "p":
            story.append(Paragraph(_esc(payload), styles["p"]))
        elif tag == "bullets":
            for item in payload:
                story.append(Paragraph(_esc(item), styles["bullet"], bulletText="•"))
            story.append(Spacer(1, 4))
        elif tag == "table":
            headers, rows = payload
            story.append(_pdf_table(headers, rows, styles))
            story.append(Spacer(1, 8))
        elif tag == "gate":
            story.append(Paragraph(_esc(payload), styles["gate"]))
        elif tag == "footer":
            story.append(Paragraph(_esc(payload), styles["footer"]))

    doc.build(story)
    return buf.getvalue()
